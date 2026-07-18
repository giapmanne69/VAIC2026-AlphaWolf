import json
import logging
import pickle
import re
import unicodedata
import math
from pathlib import Path
from typing import List, Dict, Any
from collections import Counter
import numpy as np
import pdfplumber
import docx
from openai import OpenAI

from config import settings

logger = logging.getLogger("RAGSearchTool")

# Regex dùng để tokenize tiếng Việt cho BM25
_TOKEN_PATTERN = re.compile(r"[\wÀ-ỹ]+(?:[./-][\wÀ-ỹ]+)*", re.UNICODE)

def tokenize_vietnamese(text: str) -> List[str]:
    """
    Chuẩn hóa NFC và tách từ tiếng Việt phục vụ tìm kiếm lexical.
    """
    if not text:
        return []
    normalized = unicodedata.normalize("NFC", text)
    return [token.casefold() for token in _TOKEN_PATTERN.findall(normalized)]


class RAGSearchTool:
    def __init__(self):
        self.db_path = Path(settings.FAISS_DB_DIR) / "rag_store.pkl"
        self.docs_dir = Path(settings.LEGAL_DOCS_DIR)
        
        self.client = OpenAI(
            api_key=settings.FPT_API_KEY,
            base_url=settings.FPT_BASE_URL
        )
        self.chunks: List[Dict[str, Any]] = []
        self.vectors: np.ndarray = np.array([])
        
        # Cấu hình tham số tìm kiếm lai (Hybrid Search)
        self.rrf_k = 60
        self.dense_weight = 1.0
        self.lexical_weight = 0.8
        
        # Tự động nạp cơ sở dữ liệu
        Path(settings.FAISS_DB_DIR).mkdir(parents=True, exist_ok=True)
        self.load_index()

    def load_index(self):
        """
        Nạp dữ liệu chunks và vectors từ pickle.
        """
        if self.db_path.exists():
            try:
                with open(self.db_path, "rb") as f:
                    data = pickle.load(f)
                    self.chunks = data.get("chunks", [])
                    self.vectors = np.array(data.get("vectors", []))
                logger.info(f"Đã nạp thành công RAG Database với {len(self.chunks)} phân đoạn.")
            except Exception as e:
                logger.error(f"Lỗi khi nạp RAG Database: {str(e)}")
        else:
            logger.info("Chưa có cơ sở dữ liệu RAG sẵn có. Cần chạy build_index().")

    def save_index(self):
        """
        Ghi dữ liệu chunks và vectors xuống đĩa.
        """
        try:
            with open(self.db_path, "wb") as f:
                pickle.dump({
                    "chunks": self.chunks,
                    "vectors": self.vectors.tolist()
                }, f)
            logger.info(f"Lưu cơ sở dữ liệu RAG thành công tại: {self.db_path}")
        except Exception as e:
            logger.error(f"Lỗi khi lưu RAG Database: {str(e)}")

    # ==========================================
    # PIPELINE 1 — TIỀN XỬ LÝ & XÂY DỰNG KNOWLEDGE BASE
    # ==========================================
    
    def clean_text(self, text: str) -> str:
        """
        Làm sạch nội dung: xóa khoảng trắng thừa, xóa header/footer trùng lặp,
        chuẩn hóa Unicode nhưng giữ nguyên ý nghĩa hành chính và số văn bản.
        """
        if not text:
            return ""
        # Chuẩn hóa Unicode NFC
        text = unicodedata.normalize("NFC", text)
        
        # Xóa header/footer dạng trang hoặc tiêu đề lặp đi lặp lại
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line_strip = line.strip()
            # Bỏ qua các dòng chỉ chứa số trang cô lập
            if re.match(r"^\d+$", line_strip):
                continue
            # Bỏ qua header tiêu đề lặp lại (ví dụ chế độ báo cáo)
            if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in line_strip or "Độc lập - Tự do - Hạnh phúc" in line_strip:
                continue
            cleaned_lines.append(line)
            
        text = "\n".join(cleaned_lines)
        # Xóa khoảng trắng thừa
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def analyze_structure(self, text: str) -> List[Dict[str, Any]]:
        """
        Phân tích cấu trúc hành chính Việt Nam (Điều, Khoản, Điểm, Chương)
        để phân đoạn chính xác theo quy chuẩn.
        """
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        blocks = []
        
        current_chapter = ""
        current_section = ""
        current_article = ""
        
        for p in paragraphs:
            # Nhận diện Chương
            chapter_match = re.match(r"^(CHƯƠNG\s+[IVXLCDM]+|Chương\s+\d+)\.?\s*(.*)", p, re.IGNORECASE)
            if chapter_match:
                current_chapter = p
                continue
                
            # Nhận diện Mục
            section_match = re.match(r"^(Mục\s+\d+)\.?\s*(.*)", p, re.IGNORECASE)
            if section_match:
                current_section = p
                continue
                
            # Nhận diện Điều
            article_match = re.match(r"^(Điều\s+\d+)\.?\s*(.*)", p, re.IGNORECASE)
            if article_match:
                current_article = p
                
            blocks.append({
                "content": p,
                "chapter": current_chapter,
                "section": current_section,
                "article": current_article
            })
            
        return blocks

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """
        Trích xuất văn bản từ tệp PDF sử dụng pdfplumber.
        """
        text_list = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_list.append(f"[Trang {i+1}]\n{page_text}")
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất PDF {pdf_path.name}: {str(e)}")
        return "\n\n".join(text_list)

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """
        Trích xuất văn bản từ tệp DOCX sử dụng python-docx.
        """
        text_list = []
        try:
            doc = docx.Document(docx_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    text_list.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_list.append(row_text)
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất DOCX {docx_path.name}: {str(e)}")
        return "\n".join(text_list)

    def build_index(self) -> int:
        """
        Pipeline 1: Quét tri thức, phân loại, phân tích cấu trúc, chunking, gắn metadata, embed và lưu DB.
        """
        logger.info("Khởi chạy Pipeline 1: Xây dựng Knowledge Base...")
        self.chunks = []
        vectors_list = []
        
        if not self.docs_dir.exists():
            logger.warning(f"Thư mục tri thức trống tại: {self.docs_dir}")
            return 0
            
        for filepath in self.docs_dir.rglob("*"):
            if not filepath.is_file():
                continue
            suffix = filepath.suffix.lower()
            if suffix not in [".pdf", ".docx"]:
                continue
                
            # Phân loại tài liệu dựa trên đường dẫn
            domain = "common_reporting"
            if "population" in filepath.parts or "dân cư" in str(filepath).lower():
                domain = "population"
            elif "complaint" in filepath.parts or "khiếu nại" in str(filepath).lower():
                domain = "complaints"
            elif "task" in filepath.parts or "nhiệm vụ" in str(filepath).lower():
                domain = "tasks"
                
            doc_type = "legal_document"
            if "guideline" in str(filepath).lower() or "hướng dẫn" in str(filepath).lower():
                doc_type = "reporting_guideline"
            elif "template" in str(filepath).lower() or "biểu mẫu" in str(filepath).lower():
                doc_type = "report_template"

            logger.info(f"Đọc tệp: {filepath.name} (Phân loại: Lĩnh vực={domain}, Loại={doc_type})")
            
            # Parser tài liệu
            if suffix == ".pdf":
                raw_text = self.extract_text_from_pdf(filepath)
            else:
                raw_text = self.extract_text_from_docx(filepath)
                
            # Làm sạch văn bản
            cleaned_text = self.clean_text(raw_text)
            if not cleaned_text.strip():
                continue
                
            # Phân tích cấu trúc hành chính
            structural_blocks = self.analyze_structure(cleaned_text)
            
            # Phân đoạn Chunking (khoảng 800-1000 kí tự gối đầu)
            current_chunk = []
            current_len = 0
            chunk_idx = 0
            current_pages = set()
            
            for block in structural_blocks:
                content = block["content"]
                
                # Quét số trang nếu có
                page_match = re.search(r"\[Trang (\d+)\]", content)
                if page_match:
                    current_pages.add(int(page_match.group(1)))
                    
                current_chunk.append(block)
                current_len += len(content)
                
                # Khi vượt ngưỡng chunk size, tiến hành đóng gói và gắn metadata
                if current_len >= 800:
                    chunk_content = "\n".join([b["content"] for b in current_chunk])
                    # Lưu trữ thông tin phân cấp hành chính ở chunk metadata
                    last_block = current_chunk[-1]
                    
                    self.chunks.append({
                        "chunk_id": f"{filepath.name}_chunk_{chunk_idx}",
                        "document_name": filepath.name,
                        "document_type": doc_type,
                        "domain": domain,
                        "chunk_index": chunk_idx,
                        "content": chunk_content,
                        "pages": list(current_pages),
                        "chapter": last_block["chapter"],
                        "section": last_block["section"],
                        "article": last_block["article"]
                    })
                    chunk_idx += 1
                    # Gối đầu: giữ lại 2 đoạn cuối cùng
                    current_chunk = current_chunk[-2:] if len(current_chunk) > 2 else []
                    current_len = sum([len(b["content"]) for b in current_chunk])
                    current_pages = set()
                    
            # Chunk còn sót lại cuối cùng
            if current_chunk:
                chunk_content = "\n".join([b["content"] for b in current_chunk])
                last_block = current_chunk[-1]
                self.chunks.append({
                    "chunk_id": f"{filepath.name}_chunk_{chunk_idx}",
                    "document_name": filepath.name,
                    "document_type": doc_type,
                    "domain": domain,
                    "chunk_index": chunk_idx,
                    "content": chunk_content,
                    "pages": list(current_pages),
                    "chapter": last_block["chapter"],
                    "section": last_block["section"],
                    "article": last_block["article"]
                })

        if not self.chunks:
            logger.warning("Không trích xuất được mảnh dữ liệu RAG nào.")
            return 0
            
        # Tạo embeddings qua FPT AI API
        logger.info(f"Đang sinh vector nhúng cho {len(self.chunks)} chunks bằng FPT Embedding API...")
        batch_size = 16
        for i in range(0, len(self.chunks), batch_size):
            batch = self.chunks[i:i+batch_size]
            texts = [item["content"] for item in batch]
            try:
                response = self.client.embeddings.create(
                    model=settings.EMBEDDING_MODEL,
                    input=texts
                )
                for idx, embedding_item in enumerate(response.data):
                    vectors_list.append(embedding_item.embedding)
            except Exception as e:
                logger.error(f"Lỗi gọi FPT Embeddings tại lô {i}: {str(e)}")
                dummy_vector = [0.0] * 1024
                for _ in range(len(batch)):
                    vectors_list.append(dummy_vector)
                    
        self.vectors = np.array(vectors_list, dtype=np.float32)
        self.save_index()
        return len(self.chunks)


    # ==========================================
    # PIPELINE 2 — AGENT SỬ DỤNG RAG SEARCH TOOL
    # ==========================================

    def search_dense(self, query_vector: np.ndarray, top_k: int, domain_filter: str = None) -> List[Dict[str, Any]]:
        """
        Tìm kiếm Dense Search (Cosine Similarity).
        """
        dot_products = np.dot(self.vectors, query_vector)
        norms_vectors = np.linalg.norm(self.vectors, axis=1)
        norm_query = np.linalg.norm(query_vector)
        
        norms_vectors[norms_vectors == 0] = 1e-9
        if norm_query == 0:
            norm_query = 1e-9
            
        scores = dot_products / (norms_vectors * norm_query)
        ranked_indices = np.argsort(scores)[::-1]
        
        results = []
        for idx in ranked_indices:
            chunk = self.chunks[idx]
            if domain_filter and chunk["domain"] != domain_filter:
                continue
            results.append({
                "chunk_id": chunk["chunk_id"],
                "score": float(scores[idx]),
                "chunk": chunk
            })
            if len(results) >= top_k:
                break
        return results

    def search_lexical(self, query: str, top_k: int, domain_filter: str = None) -> List[Dict[str, Any]]:
        """
        Tìm kiếm Lexical Search sử dụng giải thuật BM25 trên tập văn bản đã có.
        """
        query_tokens = tokenize_vietnamese(query)
        if not query_tokens or not self.chunks:
            return []
            
        # Lọc các chunks theo domain trước
        filtered_chunks = []
        for chunk in self.chunks:
            if domain_filter and chunk["domain"] != domain_filter:
                continue
            filtered_chunks.append(chunk)
            
        if not filtered_chunks:
            return []
            
        # Tách tokens cho từng chunk
        chunk_tokens = []
        for chunk in filtered_chunks:
            tokens = tokenize_vietnamese(" ".join((
                chunk["document_name"],
                str(chunk.get("article", "")),
                chunk["content"]
            )))
            chunk_tokens.append(tokens)
            
        # Tính toán tần suất từ DF
        doc_count = len(filtered_chunks)
        df = {}
        for tokens in chunk_tokens:
            for token in set(tokens):
                df[token] = df.get(token, 0) + 1
                
        avg_doc_len = sum(len(t) for t in chunk_tokens) / doc_count if doc_count > 0 else 1.0
        k1 = 1.5
        b = 0.75
        
        scored = []
        for idx, (chunk, tokens) in enumerate(zip(filtered_chunks, chunk_tokens)):
            tf = Counter(tokens)
            score = 0.0
            for token in query_tokens:
                if token not in tf:
                    continue
                token_df = df.get(token, 0)
                # IDF formula
                idf = math.log(1.0 + (doc_count - token_df + 0.5) / (token_df + 0.5))
                denominator = tf[token] + k1 * (1.0 - b + b * len(tokens) / avg_doc_len)
                score += idf * tf[token] * (k1 + 1.0) / denominator
                
            if score > 0:
                scored.append({
                    "chunk_id": chunk["chunk_id"],
                    "score": score,
                    "chunk": chunk
                })
                
        # Sắp xếp và trả về top_k
        scored.sort(key=lambda x: -x["score"])
        return scored[:top_k]

    def expand_parent_context(self, matched_chunks: List[Dict[str, Any]], max_tokens: int = 2048) -> List[Dict[str, Any]]:
        """
        Bảo toàn ngữ cảnh rộng hơn (Parent Context Expansion):
        Nếu tìm thấy một chunk, tự động lấy thêm 1 chunk trước và sau của cùng một tài liệu
        để tạo ngữ cảnh toàn diện, ngăn ngừa đứt đoạn thông tin luật pháp.
        """
        expanded = []
        # Nhóm chunks theo document_name
        doc_groups = {}
        for item in self.chunks:
            doc_groups.setdefault(item["document_name"], []).append(item)
            
        for matched in matched_chunks:
            chunk = matched["chunk"]
            doc_name = chunk["document_name"]
            idx = chunk["chunk_index"]
            
            all_doc_chunks = doc_groups.get(doc_name, [])
            all_doc_chunks.sort(key=lambda x: x["chunk_index"])
            
            # Lấy index trước và sau
            start_idx = max(0, idx - 1)
            end_idx = min(len(all_doc_chunks) - 1, idx + 1)
            
            # Ghép nội dung
            expanded_content = ""
            pages = set()
            articles = set()
            
            for i in range(start_idx, end_idx + 1):
                c = all_doc_chunks[i]
                expanded_content += f"\n{c['content']}"
                if c.get("pages"):
                    pages.update(c["pages"])
                if c.get("article"):
                    articles.add(c["article"])
                    
            expanded_chunk = chunk.copy()
            expanded_chunk["content"] = expanded_content.strip()
            expanded_chunk["pages"] = list(pages)
            expanded_chunk["article"] = ", ".join(filter(None, articles))
            
            expanded.append({
                "chunk_id": chunk["chunk_id"],
                "score": matched["score"],
                "chunk": expanded_chunk
            })
            
        return expanded

    def search_hybrid(self, query: str, top_k: int = 4, domain_filter: str = None) -> List[Dict[str, Any]]:
        """
        Pipeline 2 Core: Hybrid Search kết hợp Dense + Lexical (BM25) sử dụng Reciprocal Rank Fusion (RRF).
        """
        if not self.chunks or self.vectors.size == 0:
            logger.info("Index rỗng. Tiến hành build index...")
            self.build_index()
            if not self.chunks:
                return []
                
        # 1. Tính toán vector nhúng của câu hỏi để chạy Dense Search
        try:
            response = self.client.embeddings.create(
                model=settings.EMBEDDING_MODEL,
                input=[query]
            )
            query_vector = np.array(response.data[0].embedding, dtype=np.float32)
        except Exception as e:
            logger.error(f"Lỗi sinh vector nhúng câu hỏi: {str(e)}")
            return []
            
        # 2. Tìm kiếm song song 2 luồng
        dense_results = self.search_dense(query_vector, top_k=20, domain_filter=domain_filter)
        lexical_results = self.search_lexical(query, top_k=20, domain_filter=domain_filter)
        
        # 3. Kết hợp kết quả bằng Reciprocal Rank Fusion (RRF)
        rrf_scores = {}
        chunks_map = {}
        
        # Thêm điểm Dense
        for rank, item in enumerate(dense_results):
            cid = item["chunk_id"]
            chunks_map[cid] = item["chunk"]
            rrf_scores[cid] = rrf_scores.get(cid, 0.0) + self.dense_weight / (rank + self.rrf_k)
            
        # Thêm điểm Lexical
        for rank, item in enumerate(lexical_results):
            cid = item["chunk_id"]
            chunks_map[cid] = item["chunk"]
            rrf_scores[cid] = rrf_scores.get(cid, 0.0) + self.lexical_weight / (rank + self.rrf_k)
            
        # Xếp hạng lại theo điểm RRF
        sorted_rrf = sorted(rrf_scores.items(), key=lambda x: -x[1])
        
        candidates = []
        for cid, score in sorted_rrf[:top_k]:
            candidates.append({
                "chunk_id": cid,
                "score": score,
                "chunk": chunks_map[cid]
            })
            
        # 4. Mở rộng ngữ cảnh cha (Parent Context Expansion)
        expanded_candidates = self.expand_parent_context(candidates)
        return expanded_candidates

    def search(self, query: str, top_k: int = 4, domain_filter: str = None) -> List[Dict[str, Any]]:
        """
        Phương thức tương thích ngược. Trả về cấu trúc kết quả tra cứu.
        """
        raw_results = self.search_hybrid(query, top_k=top_k, domain_filter=domain_filter)
        results = []
        for item in raw_results:
            chunk = item["chunk"]
            results.append({
                "chunk_id": chunk["chunk_id"],
                "document_name": chunk["document_name"],
                "document_type": chunk["document_type"],
                "domain": chunk["domain"],
                "content": chunk["content"],
                "pages": chunk["pages"],
                "score": item["score"]
            })
        return results

    def retrieve_context(self, query: str, top_k: int = 4, domain_filter: str = None) -> str:
        """
        Ghép các đoạn tốt nhất thành Context để LLM tham chiếu trong prompt, kiểm tra trích dẫn nguồn.
        """
        results = self.search(query, top_k=top_k, domain_filter=domain_filter)
        if not results:
            return "Không tìm thấy văn bản pháp lý quy chế nào liên quan."
            
        context_parts = []
        for i, res in enumerate(results):
            pages_str = f", Trang: {res['pages']}" if res['pages'] else ""
            context_parts.append(
                f"[Tài liệu tham khảo {i+1}]: {res['document_name']} (Lĩnh vực: {res['domain']}{pages_str})\n"
                f"Nội dung quy định:\n{res['content']}\n"
                f"----------------------------------------"
            )
        return "\n\n".join(context_parts)
