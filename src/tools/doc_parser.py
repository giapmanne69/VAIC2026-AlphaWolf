import os
import logging
from pathlib import Path
import openpyxl
from docx import Document
import pdfplumber
from PIL import Image
from config import settings

try:
    import pytesseract
except ImportError:
    pytesseract = None

logger = logging.getLogger("DocParser")


class DocParser:
    def __init__(self, tesseract_cmd_path: str = None):
        logger.info("Khởi tạo DocParser để phân tích tài liệu.")
        if tesseract_cmd_path and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd_path
            logger.info(f"Đặt đường dẫn Tesseract binary: {tesseract_cmd_path}")
        elif pytesseract and os.name == "nt":
            default_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Users\DELL\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
            ]
            for p in default_paths:
                if Path(p).exists():
                    pytesseract.pytesseract.tesseract_cmd = p
                    logger.info(f"Phát hiện tự động Tesseract binary tại: {p}")
                    break

    def parse(self, file_path: str) -> str:
        path = Path(file_path)
        logger.info(f"Đang phân tích tệp tin đầu vào: {path.name}")
        if not path.exists():
            logger.error(f"Tệp tin không tồn tại: {file_path}")
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
            
        ext = path.suffix.lower()
        result_text = ""
        if ext == ".docx":
            result_text = self.parse_docx(path)
        elif ext in [".xlsx", ".xls"]:
            result_text = self.parse_xlsx(path)
        elif ext == ".pdf":
            result_text = self.parse_pdf(path)
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            result_text = self.parse_image(path)
        else:
            logger.error(f"Định dạng tệp không được hỗ trợ: {ext}")
            raise ValueError(f"Định dạng tệp không được hỗ trợ: {ext}")
            
        logger.info(f"Hoàn thành phân tích {path.name}. Số ký tự trích xuất: {len(result_text)}")
        return result_text

    def parse_docx(self, path: Path) -> str:
        logger.info(f"Đang đọc nội dung tệp Word (.docx): {path.name}")
        doc = Document(path)
        content = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                p_text = element.text if hasattr(element, 'text') else ""
                for p in doc.paragraphs:
                    if p._element == element:
                        p_text = p.text
                        break
                if p_text.strip():
                    content.append(p_text.strip())
            elif element.tag.endswith('tbl'):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    table_content = []
                    for row in table.rows:
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        table_content.append(" | ".join(row_cells))
                    content.append("\n[BẢNG BIỂU]:\n" + "\n".join(table_content) + "\n")
                    
        return "\n".join(content)

    def parse_xlsx(self, path: Path) -> str:
        logger.info(f"Đang đọc bảng tính Excel (.xlsx): {path.name}")
        wb = openpyxl.load_workbook(path, data_only=True)
        content = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            logger.info(f"Đọc trang tính: {sheet_name} (Số dòng tối đa: {sheet.max_row})")
            content.append(f"\n--- Trang tính: {sheet_name} ---")
            
            for r in range(1, sheet.max_row + 1):
                row_vals = []
                has_data = False
                for c in range(1, sheet.max_column + 1):
                    val = sheet.cell(row=r, column=c).value
                    if val is not None:
                        has_data = True
                        row_vals.append(str(val).strip())
                    else:
                        row_vals.append("")
                if has_data:
                    content.append(" | ".join(row_vals))
                    
        return "\n".join(content)

    def parse_pdf(self, path: Path) -> str:
        logger.info(f"Đang đọc tệp PDF: {path.name}")
        content = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    logger.info(f"Đã đọc văn bản từ trang PDF {i+1}")
                    content.append(f"\n--- Trang PDF {i+1} ---")
                    content.append(text.strip())
                else:
                    logger.warning(f"Trang PDF {i+1} không có văn bản. Thử chạy OCR...")
                    content.append(f"\n--- Trang PDF Scan {i+1} ---")
                    content.append("[Cảnh báo: Phát hiện trang quét PDF Scan dạng ảnh, cần chạy OCR]")
                        
        return "\n".join(content)

    def parse_image(self, path: Path) -> str:
        logger.info(f"Đang chạy OCR hình ảnh. Thử nghiệm gọi FPT Vision Model: {settings.VISION_MODEL}")
        try:
            import base64
            from openai import OpenAI
            
            with open(path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode("utf-8")
                
            client = OpenAI(api_key=settings.FPT_API_KEY, base_url=settings.FPT_BASE_URL)
            response = client.chat.completions.create(
                model=settings.VISION_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text", 
                                "text": "Hãy chuyển tải toàn bộ văn bản tiếng Việt có trong hình ảnh này thành định dạng text, giữ nguyên cấu trúc dòng và bảng số liệu nếu có. Không giải thích thêm."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2048,
                temperature=0.1
            )
            text_result = response.choices[0].message.content
            logger.info("Hoàn tất trích xuất văn bản từ ảnh qua Vision Model.")
            return text_result.strip()
        except Exception as api_err:
            logger.warning(f"Không thể gọi FPT Vision Model ({str(api_err)}). Thử fallback sang Tesseract local...")
            
            if not pytesseract:
                logger.error("Thư viện pytesseract chưa được cài đặt.")
                return f"[Lỗi trích xuất ảnh: Thư viện pytesseract chưa được cài đặt và API Vision gặp lỗi: {str(api_err)}]"
                
            try:
                image = Image.open(path)
                text = pytesseract.image_to_string(image, lang="vie+eng")
                logger.info("Hoàn tất chạy OCR hình ảnh bằng Tesseract local.")
                return text.strip()
            except Exception as e:
                logger.exception(f"Lỗi OCR hình ảnh {path.name}:")
                return f"[Lỗi OCR ảnh {path.name}: {str(e)}]"
