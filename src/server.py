import os
import uuid
import json
import logging
import shutil
import io
import re
from pathlib import Path
from typing import List
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document

from config import settings
from src.agent import AgenticReportAgent

# Thiết lập log cho Backend
logger = logging.getLogger("FastAPIBackend")

app = FastAPI(
    title="VAIC AI Report Agent API",
    description="Backend API hỗ trợ Tác tử AI phân tích và tự động điền báo cáo hành chính",
    version="1.0"
)

# Cấu hình CORS để React Frontend gọi API thuận tiện
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Trong thực tế sản xuất cần giới hạn domain cụ thể
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Thư mục tạm thời lưu trữ session tệp mẫu trống
SESSIONS_DIR = Path(settings.DATA_DIR) / "sessions"
SESSIONS_DIR.mkdir(parents=True, exist_ok=True)


class RenderRequest(BaseModel):
    session_id: str
    kpi_data: dict
    remarks: str


def clean_session_dir(session_id: str):
    """
    Dọn dẹp thư mục tạm của session sau khi hoàn thành.
    """
    session_path = SESSIONS_DIR / session_id
    if session_path.exists():
        try:
            shutil.rmtree(session_path)
            logger.info(f"Đã dọn dẹp thư mục session: {session_id}")
        except Exception as e:
            logger.error(f"Không thể dọn dẹp session {session_id}: {str(e)}")


def replace_text_in_paragraph(p, target, replacement):
    """
    Thay thế chữ trong đoạn văn mà không làm hỏng định dạng (bold, italic, font).
    Đặc biệt, nếu đoạn văn hoặc run có chứa thẻ 'w:drawing' (ảnh/logo/chữ ký),
    chỉ thay thế chữ ở các run không có ảnh để tránh làm mất ảnh trong tài liệu.
    """
    if target not in p.text:
        return
        
    xml_str = p._element.xml
    has_image = "w:drawing" in xml_str or "w:pict" in xml_str
    
    if has_image:
        # Đoạn văn có ảnh: Chỉ duyệt thay thế trên những run không có ảnh
        for run in p.runs:
            run_xml = run.element.xml
            if "w:drawing" not in run_xml and "w:pict" not in run_xml:
                if target in run.text:
                    run.text = run.text.replace(target, replacement)
        return

    # Đoạn văn thường (không có ảnh): Thay thế bảo toàn định dạng
    if len(p.runs) == 1:
        p.runs[0].text = p.runs[0].text.replace(target, replacement)
    else:
        # Ghép text đã thay thế và gán vào run đầu tiên, xóa các run sau để giữ định dạng
        full_text = p.text.replace(target, replacement)
        p.runs[0].text = full_text
        for r in p.runs[1:]:
            r.text = ""


def auto_inject_jinja_tags(template_bytes: bytes) -> bytes:
    """
    Tự động phân tích cấu trúc tệp Word mẫu trống (.docx) do người dùng tải lên,
    tìm các ký tự giữ chỗ cổ điển (như [..........], [Điền], AI Assistant Template)
    và thay thế bằng thẻ Jinja {{ ... }} tương ứng để docxtpl có thể điền được.
    """
    try:
        doc = Document(io.BytesIO(template_bytes))
        
        # 1. Quét qua các đoạn văn (paragraphs) để chèn các thẻ nhận xét theo mẫu chung
        for p in doc.paragraphs:
            text = p.text
            if "[Kỳ báo cáo]" in text:
                replace_text_in_paragraph(p, "[Kỳ báo cáo]", "{{ so_ngay_thang_nam }}")
            if "[Kỳ tiếp theo]" in text:
                replace_text_in_paragraph(p, "[Kỳ tiếp theo]", "Kỳ tiếp theo")
            if "⚡ [AI Assistant Template]" in text or "AI Assistant Template" in text:
                target_str = "⚡ [AI Assistant Template]" if "⚡ [AI Assistant Template]" in text else "AI Assistant Template"
                replace_text_in_paragraph(p, target_str, "{{ nhan_xet_ai }}")
            
        # 2. Quét qua các bảng biểu (tables)
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                
                # Kiểm tra thông tin hành chính ở cột đầu tiên
                first_cell = cells[0]
                first_cell_text = first_cell.text.strip()
                if "UBND QUẬN/HUYỆN" in first_cell_text:
                    for p in first_cell.paragraphs:
                        replace_text_in_paragraph(p, "[..........]", "{{ ten_co_quan_cap_on }}")
                        replace_text_in_paragraph(p, "[..........]", "{{ ten_co_quan_cap_duoi }}")
                    
                    # Địa danh ngày tháng năm
                    if len(cells) > 1:
                        cell_1 = cells[1]
                        for p in cell_1.paragraphs:
                            for text_to_replace in ["ngày ...... tháng ...... năm 20...", "ngày ...... tháng"]:
                                if text_to_replace in p.text:
                                    replace_text_in_paragraph(p, text_to_replace, "{{ so_ngay_thang_nam }}")
                        
                # Tên chỉ tiêu ở cột index 1
                indicator_name = cells[1].text.strip()
                indicator_lower = indicator_name.lower()
                
                # Ký tên chủ tịch ở cuối bảng
                last_cell = cells[-1]
                last_cell_text = last_cell.text.strip()
                if "CHỦ TỊCH" in last_cell_text or "Chủ tịch" in last_cell_text:
                    for p in last_cell.paragraphs:
                        replace_text_in_paragraph(p, "[Điền họ và tên Chủ tịch]", "{{ ten_chu_tich }}")
                
                # So khớp từ khóa để map chỉ tiêu
                var_base = None
                if "thu ngân sách" in indicator_lower:
                    var_base = "tong_thu_ngan_sach_nha_nuoc"
                elif "chi ngân sách" in indicator_lower:
                    var_base = "tong_chi_ngan_sach_dia_phuong"
                elif "khai sinh" in indicator_lower:
                    var_base = "dang_ky_khai_sinh"
                elif "khai tử" in indicator_lower or "khai tu" in indicator_lower:
                    var_base = "dang_ky_khai_tu"
                elif "cư trú" in indicator_lower or "tạm trú" in indicator_lower:
                    var_base = "tam_tru_moi"
                elif "chứng thực" in indicator_lower:
                    var_base = "chung_thuc_chu_ky"
                elif "an ninh" in indicator_lower or "trật tự" in indicator_lower:
                    var_base = "vi_pham_an_ninh_trat_tu"
                
                if var_base:
                    # Kỳ trước (Cột 4 / Index 3)
                    if len(cells) > 3:
                        cell_3 = cells[3]
                        for p in cell_3.paragraphs:
                            if "[Điền]" in p.text:
                                replace_text_in_paragraph(p, "[Điền]", f"{{{{ {var_base}_ky_truoc }}}}")
                            elif p.text.strip() == "":
                                if not p.runs:
                                    p.add_run(f"{{{{ {var_base}_ky_truoc }}}}")
                                else:
                                    p.runs[0].text = f"{{{{ {var_base}_ky_truoc }}}}"
                                    
                    # Kỳ báo cáo (Cột 5 / Index 4)
                    if len(cells) > 4:
                        cell_4 = cells[4]
                        for p in cell_4.paragraphs:
                            if "[Điền]" in p.text:
                                replace_text_in_paragraph(p, "[Điền]", f"{{{{ {var_base}_ky_bao_cao }}}}")
                            elif p.text.strip() == "":
                                if not p.runs:
                                    p.add_run(f"{{{{ {var_base}_ky_bao_cao }}}}")
                                else:
                                    p.runs[0].text = f"{{{{ {var_base}_ky_bao_cao }}}}"
                        
        output_stream = io.BytesIO()
        doc.save(output_stream)
        logger.info("Tự động chèn thẻ Jinja thành công vào tệp biểu mẫu.")
        return output_stream.getvalue()
    except Exception as e:
        logger.exception("Lỗi khi tự động chèn thẻ Jinja vào biểu mẫu trống:")
        raise e


@app.get("/api/health")
def health_check():
    return {"status": "healthy", "model": settings.LLM_MODEL}


@app.get("/api/agent/style")
def get_style_preferences():
    """
    Lấy thói quen văn phong từ Long-term Memory.
    """
    try:
        agent = AgenticReportAgent()
        preferences = agent.long_memory.get_style_preferences()
        return {"preferences": preferences}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/agent/style")
def add_style_preference(key: str = Form(...), val: str = Form(...)):
    """
    Lưu thói quen văn phong mới vào Long-term Memory.
    """
    try:
        agent = AgenticReportAgent()
        agent.long_memory.add_style_preference(key, val)
        return {"status": "success", "message": f"Đã lưu văn phong '{key}': '{val}'"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/agent/run")
async def run_agent(
    template: UploadFile = File(...),
    raws: List[UploadFile] = File(...),
    fpt_api_key: str = Form(None)
):
    """
    Khởi chạy Tác tử AI theo mô hình ReAct Loop và Stream logs tư duy thời gian thực về React Client.
    """
    # Khởi tạo session_id độc bản
    session_id = str(uuid.uuid4())
    session_path = SESSIONS_DIR / session_id
    session_path.mkdir(parents=True, exist_ok=True)

    # Lưu file biểu mẫu trống (tự động chèn thẻ Jinja)
    template_ext = Path(template.filename).suffix
    template_bytes = await template.read()
    if template_ext.lower() == ".docx":
        try:
            template_bytes = auto_inject_jinja_tags(template_bytes)
            logger.info("Đã tự động chèn thẻ Jinja thành công vào biểu mẫu trống.")
        except Exception as e:
            logger.error(f"Lỗi khi tự động chèn thẻ Jinja: {str(e)}")
            
    template_save_path = session_path / f"template{template_ext}"
    with open(template_save_path, "wb") as f:
        f.write(template_bytes)

    # Lưu danh sách báo cáo phòng ban thô
    raw_save_paths = []
    for raw_file in raws:
        raw_ext = Path(raw_file.filename).suffix
        raw_name = Path(raw_file.filename).stem
        # Đảm bảo không trùng tên file
        raw_save_path = session_path / f"{raw_name}_{uuid.uuid4().hex[:6]}{raw_ext}"
        with open(raw_save_path, "wb") as f:
            f.write(await raw_file.read())
        raw_save_paths.append(str(raw_save_path))

    # Cấu hình API Key tạm thời nếu cán bộ nhập tay ở Client
    if fpt_api_key:
        os.environ["FPT_API_KEY"] = fpt_api_key

    # Khởi tạo Agent
    try:
        agent = AgenticReportAgent()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Không thể khởi tạo Agent: {str(e)}")

    output_report_path = session_path / "output_bao_cao.docx"

    def event_generator():
        # Gửi sự kiện mở đầu chứa session_id
        yield f"data: {json.dumps({'status': 'init', 'session_id': session_id})}\n\n"

        # Khởi chạy ReAct loop qua generator
        react_gen = agent.run_react_agent_generator(
            template_path=str(template_save_path),
            raw_paths=raw_save_paths,
            output_path=str(output_report_path),
            session_id=session_id
        )

        try:
            for step_data in react_gen:
                yield f"data: {json.dumps(step_data, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("Lỗi trong event_generator SSE:")
            error_payload = {
                'status': 'error',
                'message': f'Lỗi phía server khi xử lý luồng SSE: {str(e)}'
            }
            yield f"data: {json.dumps(error_payload, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.post("/api/agent/render-docx")
def render_docx(req: RenderRequest, background_tasks: BackgroundTasks):
    """
    Nhận số liệu đã cán bộ hiệu chỉnh trên UI (Human-in-the-loop),
    tiến hành điền mẫu Word cuối cùng và trả tệp tin về trình duyệt.
    Sau khi trả tệp tin, dọn dẹp thư mục tạm session ở nền.
    """
    session_id = req.session_id
    session_path = SESSIONS_DIR / session_id
    
    # Tìm tệp biểu mẫu trống của session
    template_files = list(session_path.glob("template.*"))
    if not template_files:
        raise HTTPException(status_code=404, detail="Không tìm thấy tệp biểu mẫu mẫu cho session này.")
    
    template_path = template_files[0]
    output_docx_path = session_path / "Bao_cao_tong_hop.docx"

    try:
        agent = AgenticReportAgent()
        
        remarks_dict = {}
        remarks = req.remarks or ""
        if isinstance(remarks, dict):
            remarks_dict = remarks
        else:
            remarks_dict["nhan_xet_ai"] = remarks.strip()

        # Điền biểu mẫu thông qua công cụ
        agent.execute_agent_tool("render_docx_report_tool", {
            "template_path": str(template_path),
            "kpi_data": req.kpi_data,
            "remarks_dict": remarks_dict,
            "output_path": str(output_docx_path)
        })

        return FileResponse(
            path=output_docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"Bao_cao_tong_hop_{session_id[:8]}.docx"
        )
    except Exception as e:
        logger.exception("Lỗi khi điền mẫu và kết xuất báo cáo cuối cùng:")
        raise HTTPException(status_code=500, detail=str(e))


# Tích hợp phục vụ giao diện tĩnh (Frontend) từ thư mục src/static
from fastapi.staticfiles import StaticFiles
static_path = Path(__file__).resolve().parent / "static"
static_path.mkdir(parents=True, exist_ok=True)
app.mount("/", StaticFiles(directory=str(static_path), html=True), name="static")

