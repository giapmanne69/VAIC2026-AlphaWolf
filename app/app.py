import sys
from pathlib import Path

# Tự động cấu hình PYTHONPATH trỏ về thư mục gốc của dự án
sys.path.append(str(Path(__file__).resolve().parent.parent))

import streamlit as st
import json
import tempfile
import pandas as pd
import io
import logging
from docxtpl import DocxTemplate
from docx import Document

# Khởi tạo logger cho Frontend
logger = logging.getLogger("StreamlitApp")

# Thiết lập cấu hình trang Streamlit (phải ở dòng đầu tiên)
st.set_page_config(
    page_title="Hệ thống Tác tử Tổng hợp Báo cáo UBND Phường",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS để làm nổi bật thẩm mỹ (Gradients, Glassmorphism, Google Fonts)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
    }
    
    /* Cấu hình tiêu đề gradient */
    .title-gradient {
        background: linear-gradient(135deg, #FF4B4B, #FF8F00, #00C853);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        font-size: 2.8rem;
        margin-bottom: 0.5rem;
    }
    
    .subtitle {
        color: #757575;
        font-size: 1.1rem;
        margin-bottom: 2rem;
    }
    
    /* Thiết kế Glassmorphism cho các khung thông tin */
    .glass-card {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 12px;
        padding: 20px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.2);
        backdrop-filter: blur(4px);
        -webkit-backdrop-filter: blur(4px);
        margin-bottom: 1.5rem;
    }
    
    .status-badge-pass {
        background-color: #E8F5E9;
        color: #2E7D32;
        padding: 4px 12px;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.85rem;
    }
    
    .status-badge-fail {
        background-color: #FFEBEE;
        color: #C62828;
        padding: 4px 12px;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.85rem;
    }
</style>
""", unsafe_allow_html=True)

# Lớp Agent và bộ nhớ
from src.agent import AgenticReportAgent
from config import settings

# Khởi tạo hoặc tái sử dụng instance Agent trong Session State của Streamlit
if "agent" not in st.session_state:
    st.session_state.agent = AgenticReportAgent()

agent = st.session_state.agent


# --- HELPER: TÁCH NHẬN XÉT THEO KHỐI ---
def parse_combined_remarks(combined_text: str) -> dict:
    """
    Tách chuỗi văn bản gộp từ UI thành các phần nhận xét riêng lẻ dựa trên tiêu đề khối.
    """
    parts = {
        "nhan_xet_ai_kinh_te": "",
        "nhan_xet_ai_van_hoa_xa_hoi": "",
        "nhan_xet_ai_quoc_phong_an_ninh": "",
        "nhan_xet_ai_phuong_huong": ""
    }
    
    import re
    # Khối Kinh tế
    match_kt = re.search(r'===\s*KHỐI KINH TẾ\s*===\n(.*?)(?===?\s*KHỐI VĂN HÓA|==?\s*KHỐI QUỐC PHÒNG|==?\s*PHƯƠNG HƯỚNG|$)', combined_text, re.DOTALL | re.IGNORECASE)
    if match_kt:
        parts["nhan_xet_ai_kinh_te"] = match_kt.group(1).strip()
        
    # Khối Văn hóa - Xã hội
    match_vh = re.search(r'===\s*KHỐI VĂN HÓA\s*-\s*XÃ HỘI\s*===\n(.*?)(?===?\s*KHỐI QUỐC PHÒNG|==?\s*PHƯƠNG HƯỚNG|$)', combined_text, re.DOTALL | re.IGNORECASE)
    if match_vh:
        parts["nhan_xet_ai_van_hoa_xa_hoi"] = match_vh.group(1).strip()
        
    # Khối Quốc phòng - An ninh
    match_qp = re.search(r'===\s*KHỐI QUỐC PHÒNG\s*-\s*AN NINH\s*===\n(.*?)(?===?\s*PHƯƠNG HƯỚNG|$)', combined_text, re.DOTALL | re.IGNORECASE)
    if match_qp:
        parts["nhan_xet_ai_quoc_phong_an_ninh"] = match_qp.group(1).strip()
        
    # Phương hướng
    match_ph = re.search(r'===\s*PHƯƠNG HƯỚNG KỲ TỚI\s*===\n(.*)', combined_text, re.DOTALL | re.IGNORECASE)
    if match_ph:
        parts["nhan_xet_ai_phuong_huong"] = match_ph.group(1).strip()
        
    return parts


# --- HELPER 1: TỰ ĐỘNG CHÈN THẺ JINJA CHO BIỂU MẪU GỐC CỦA CÁN BỘ ---
def auto_inject_jinja_tags(template_bytes: bytes) -> bytes:
    """
    Tự động phân tích cấu trúc tệp Word mẫu trống (.docx) do người dùng tải lên,
    tìm các ký tự giữ chỗ cổ điển (như [..........], [Điền], AI Assistant Template)
    và thay thế bằng thẻ Jinja {{ ... }} tương ứng để docxtpl có thể điền được.
    """
    try:
        doc = Document(io.BytesIO(template_bytes))
        
        # 1. Quét qua các đoạn văn (paragraphs) để chèn các thẻ nhận xét riêng biệt
        ai_template_counter = 0
        for p in doc.paragraphs:
            text = p.text
            if "[Kỳ báo cáo]" in text:
                text = text.replace("[Kỳ báo cáo]", "{{ so_ngay_thang_nam }}")
            if "[Kỳ tiếp theo]" in text:
                text = text.replace("[Kỳ tiếp theo]", "Kỳ tiếp theo")
            if "⚡ [AI Assistant Template]" in text or "AI Assistant Template" in text:
                ai_template_counter += 1
                if ai_template_counter == 1:
                    text = "{{ nhan_xet_ai_kinh_te }}"
                elif ai_template_counter == 2:
                    text = "{{ nhan_xet_ai_van_hoa_xa_hoi }}"
                elif ai_template_counter == 3:
                    text = "{{ nhan_xet_ai_quoc_phong_an_ninh }}"
                else:
                    text = "{{ nhan_xet_ai_phuong_huong }}"
            p.text = text
            
        # 2. Quét qua các bảng biểu (tables)
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                
                # Kiểm tra thông tin hành chính ở cột đầu tiên
                first_cell_text = cells[0].text.strip()
                if "UBND QUẬN/HUYỆN" in first_cell_text:
                    text = cells[0].text
                    if "[..........]" in text:
                        text = text.replace("[..........]", "{{ ten_co_quan_cap_on }}", 1)
                    if "[..........]" in text:
                        text = text.replace("[..........]", "{{ ten_co_quan_cap_duoi }}", 1)
                    cells[0].text = text
                    
                    # Địa danh ngày tháng năm
                    if len(cells) > 1 and "ngày ...... tháng" in cells[1].text:
                        cells[1].text = "{{ so_ngay_thang_nam }}"
                        
                # Tên chỉ tiêu ở cột index 1
                indicator_name = cells[1].text.strip()
                indicator_lower = indicator_name.lower()
                
                # Ký tên chủ tịch ở cuối bảng
                last_cell_text = cells[-1].text.strip()
                if "CHỦ TỊCH" in last_cell_text or "Chủ tịch" in last_cell_text:
                    text = cells[-1].text
                    if "[Điền họ và tên Chủ tịch]" in text:
                        cells[-1].text = text.replace("[Điền họ và tên Chủ tịch]", "{{ ten_chu_tich }}")
                
                # So khớp từ khóa để map chỉ tiêu
                var_base = None
                if "thu ngân sách" in indicator_lower:
                    var_base = "tong_thu_ngan_sach_nha_nuoc"
                elif "chi ngân sách" in indicator_lower:
                    var_base = "tong_chi_ngan_sach_dia_phuong"
                elif "khai sinh" in indicator_lower:
                    var_base = "dang_ky_khai_sinh"
                elif "khai tử" in indicator_lower or "khai tu" in indicator_lower:
                    var_base = "dang_ky_khai_tu"
                elif "cư trú" in indicator_lower or "tạm trú" in indicator_lower:
                    var_base = "tam_tru_moi"
                elif "chứng thực" in indicator_lower:
                    var_base = "chung_thuc_chu_ky"
                elif "an ninh" in indicator_lower or "trật tự" in indicator_lower:
                    var_base = "vi_pham_an_ninh_trat_tu"
                
                if var_base:
                    # Kỳ trước (Cột 4 / Index 3)
                    if len(cells) > 3 and ("[Điền]" in cells[3].text or cells[3].text.strip() == ""):
                        cells[3].text = f"{{{{ {var_base}_ky_truoc }}}}"
                    # Kỳ báo cáo (Cột 5 / Index 4)
                    if len(cells) > 4 and ("[Điền]" in cells[4].text or cells[4].text.strip() == ""):
                        cells[4].text = f"{{{{ {var_base}_ky_bao_cao }}}}"
                        
        output_stream = io.BytesIO()
        doc.save(output_stream)
        logger.info("Tự động chèn thẻ Jinja thành công vào tệp biểu mẫu.")
        return output_stream.getvalue()
    except Exception as e:
        logger.exception("Lỗi khi tự động chèn thẻ Jinja vào biểu mẫu trống:")
        raise e


# --- HELPER 2: RENDER WORD IN-MEMORY ---
def render_docx_in_memory(template_bytes: bytes, kpi_data: dict, remarks: str) -> bytes:
    """
    Render biểu mẫu Word trực tiếp trên bộ nhớ RAM để đảm bảo tính đồng bộ dữ liệu
    và cho phép cập nhật tức thì các chỉnh sửa thủ công của cán bộ (Human-in-the-loop).
    """
    try:
        doc = DocxTemplate(io.BytesIO(template_bytes))
        context = kpi_data.copy()
        
        # Xử lý dọn sạch các giá trị None/null tránh in chữ None ra Word
        for k, v in list(context.items()):
            if v is None or v == "None" or v == "null":
                context[k] = ""
                
        # Phân tách nhận xét tổng hợp từ UI ra các trường riêng lẻ
        remarks_parts = parse_combined_remarks(remarks)
        context.update(remarks_parts)
        context["nhan_xet_ai"] = remarks # Giữ cho tương thích ngược
        
        logger.info(f"Đang render Word trực tiếp trên bộ nhớ với {len(kpi_data)} chỉ số.")
        doc.render(context)
        
        output_stream = io.BytesIO()
        doc.save(output_stream)
        return output_stream.getvalue()
    except Exception as e:
        logger.exception("Lỗi khi render biểu mẫu Word trong bộ nhớ:")
        raise e


# --- SIDEBAR: CẤU HÌNH HỆ THỐNG ---
st.sidebar.markdown("<h2 style='font-weight:800;'>⚙️ CẤU HÌNH AI</h2>", unsafe_allow_html=True)
st.sidebar.info("Hệ thống đang cấu hình chạy trực tiếp với mô hình Llama-3.3-70B-Instruct thông qua FPT AI Factory API.")

# Điền nhanh API Key
api_key = st.sidebar.text_input("FPT AI Factory API Key", value=settings.FPT_API_KEY, type="password")
model_name = st.sidebar.text_input("LLM Model Name", value=settings.LLM_MODEL)

if api_key != settings.FPT_API_KEY:
    settings.FPT_API_KEY = api_key
    agent.client.api_key = api_key
    logger.info("Cập nhật FPT_API_KEY từ sidebar.")

# --- MÀN HÌNH CHÍNH ---
st.markdown("<h1 class='title-gradient'>🏛️ UBND Phường - Agentic AI Report Hub</h1>", unsafe_allow_html=True)
st.markdown("<p class='subtitle'>Tác tử AI hỗ trợ tự động hóa phân tích biểu mẫu động, trích xuất dữ liệu đa nguồn và viết báo cáo hành chính công chuẩn quy định.</p>", unsafe_allow_html=True)

col_upload, col_process = st.columns([1, 2])

with col_upload:
    st.markdown("### 📥 Tải dữ liệu đầu vào")
    
    # 1. Upload file mẫu Word trống
    template_file = st.file_uploader(
        "Tải lên Biểu mẫu trống (.docx)", 
        type=["docx"], 
        help="Cán bộ tải biểu mẫu trống của phường có chứa ngoặc vuông hoặc dấu gạch dưới để AI tự động điền."
    )
    
    # 2. Upload các file báo cáo phòng ban chuyên ngành thô
    raw_files = st.file_uploader(
        "Tải lên các Báo cáo thô phòng ban", 
        type=["docx", "xlsx", "pdf", "png", "jpg"], 
        accept_multiple_files=True,
        help="Tải lên báo cáo thô của Tư pháp, Một cửa, Địa chính, Công an... định dạng Word, Excel, PDF hoặc ảnh chụp."
    )
    
    # 3. Custom RAG Context input
    st.markdown("---")
    st.markdown("#### 📖 Tri thức luật bổ sung (RAG)")
    rag_context = st.text_area(
        "Cơ sở pháp lý cần viện dẫn (RAG):", 
        value="Nghị định 61/2018/NĐ-CP ngày 23/4/2018 của Chính phủ quy định về thực hiện cơ chế một cửa, một cửa liên thông trong giải quyết thủ tục hành chính.",
        height=100
    )

with col_process:
    st.markdown("### ⚙️ Tiến trình trích xuất & Sinh báo cáo")
    
    if not template_file or not raw_files:
        st.warning("Vui lòng tải lên đầy đủ file biểu mẫu trống và ít nhất 1 báo cáo thô để bắt đầu.")
    else:
        if st.button("🚀 BẮT ĐẦU XỬ LÝ REPORT PIPELINE", type="primary", width="stretch"):
            logger.info("--- BẮT ĐẦU CHẠY PIPELINE TRÍCH XUẤT ---")
            
            # Lưu trữ file bytes gốc và chuẩn bị tên file đầu ra
            original_bytes = template_file.getvalue()
            st.session_state.output_file_name = f"Bao_cao_tong_hop_{Path(template_file.name).stem}.docx"
            
            # Sử dụng thư mục tạm thời để lưu file vật lý chạy Parser
            with tempfile.TemporaryDirectory() as tmp_dir:
                tmp_dir_path = Path(tmp_dir)
                
                temp_template_path = tmp_dir_path / template_file.name
                with open(temp_template_path, "wb") as f:
                    f.write(original_bytes)
                
                temp_raw_paths = []
                for rf in raw_files:
                    path = tmp_dir_path / rf.name
                    with open(path, "wb") as f:
                        f.write(rf.getbuffer())
                    temp_raw_paths.append(str(path))
                
                try:
                    # --- STAGE 1: Phân tích biểu mẫu mẫu ---
                    with st.status("🔍 Stage 1: Đang quét cấu hình biểu mẫu mẫu trống...", expanded=True) as status:
                        schema = agent.parse_template_schema(str(temp_template_path))
                        st.write("Schema chỉ tiêu phát hiện được từ mẫu thô:")
                        st.json(schema)
                        status.update(label="Stage 1: Hoàn tất phân tích mẫu!", state="complete")
                    
                    # Tự động chèn thẻ Jinja vào biểu mẫu gốc trước khi render trên bộ nhớ
                    with st.status("⚡ Đang tự động gắn thẻ liên kết dữ liệu vào mẫu trống...", expanded=True) as status:
                        st.session_state.template_bytes = auto_inject_jinja_tags(original_bytes)
                        status.update(label="Đã tự động liên kết thẻ mẫu thành công!", state="complete")
                    
                    # --- STAGE 2 & 3: Trích xuất số liệu ---
                    with st.status("🛡️ Stage 2 & 3: Đang ẩn danh PII và trích xuất số liệu...", expanded=True) as status:
                        raw_data = agent.extract_from_raw_inputs(schema, temp_raw_paths)
                        st.write("Dữ liệu trích xuất:")
                        st.json(raw_data)
                        status.update(label="Stage 2 & 3: Hoàn tất trích xuất!", state="complete")
                    
                    # --- STAGE 4: Kiểm chéo và tự sửa lỗi ---
                    with st.status("🧮 Stage 4: Đang kiểm tra logic số liệu chéo...", expanded=True) as status:
                        final_data, failures = agent.run_validation_and_self_correction(raw_data)
                        
                        if not failures:
                            st.markdown("<span class='status-badge-pass'>PASS</span> Không phát hiện lỗi logic.", unsafe_allow_html=True)
                        else:
                            for fail in failures:
                                status_badge = "FAIL (ERROR)" if fail["severity"] == "error" else "WARNING"
                                badge_class = "status-badge-fail" if fail["severity"] == "error" else "status-badge-pass"
                                st.markdown(
                                    f"- <span class='{badge_class}'>{status_badge}</span> **{fail['id']}**: {fail['description']}<br>"
                                    f"  *Công thức: `{fail['formula']}` ({fail['error_msg']})*", 
                                    unsafe_allow_html=True
                                )
                        status.update(label="Stage 4: Kiểm lỗi hoàn tất!", state="complete")
                    
                    # --- STAGE 6: Sinh nhận định ---
                    with st.status("📝 Stage 6: Đang sinh đoạn văn nhận định...", expanded=True) as status:
                        remarks = agent.generate_final_report(
                            kpi_data=final_data,
                            template_path=str(temp_template_path),
                            output_path=str(tmp_dir_path / "mock.docx"),
                            rag_context=rag_context
                        )
                        st.write("Nhận định từ LLM:")
                        st.info(remarks)
                        
                        # Lưu trữ kết quả trích xuất vào Session State
                        st.session_state.final_remarks = remarks
                        st.session_state.final_kpi = final_data
                        
                        status.update(label="Stage 6: Hoàn tất sinh nhận định!", state="complete")
                        
                    st.balloons()
                    logger.info("Pipeline trích xuất báo cáo chạy thành công.")
                except Exception as e:
                    logger.exception("Lỗi xảy ra trong quá trình chạy Report Pipeline:")
                    st.error(f"Đã xảy ra lỗi nghiêm trọng: {str(e)}. Xem chi tiết tại data/agent_execution.log")

        # Màn hình kết quả khi dữ liệu đã sẵn sàng trong session state
        if "final_kpi" in st.session_state and "template_bytes" in st.session_state:
            st.markdown("---")
            st.markdown("### 🏆 Báo cáo đã sẵn sàng tải về")
            
            # Khung Human-in-the-loop cho phép sửa trực tiếp nhận xét
            st.markdown("#### 🖊️ Hiệu chỉnh nhận xét tự động (Human-in-the-loop)")
            user_edited_remarks = st.text_area(
                "Bạn có thể tinh chỉnh đoạn văn dưới đây. File Word tải xuống sẽ tự động cập nhật theo nội dung mới này:",
                value=st.session_state.final_remarks,
                height=150
            )
            
            # Đồng bộ chỉnh sửa vào bộ nhớ dài hạn
            if user_edited_remarks != st.session_state.final_remarks:
                agent.long_memory.add_style_preference("user_adjusted_remarks", user_edited_remarks)
                st.session_state.final_remarks = user_edited_remarks
                st.toast("Đã lưu thói quen chỉnh sửa và cập nhật tệp tải về!", icon="🧠")
            
            # RENDER TRÊN BỘ NHỚ RAM TRƯỚC KHI TẢI VỀ
            try:
                docx_bytes = render_docx_in_memory(
                    template_bytes=st.session_state.template_bytes,
                    kpi_data=st.session_state.final_kpi,
                    remarks=st.session_state.final_remarks
                )
                
                st.download_button(
                    label="📥 Tải xuống báo cáo hoàn chỉnh (.docx)",
                    data=docx_bytes,
                    file_name=st.session_state.output_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width="stretch"
                )
            except Exception as e:
                st.error(f"Lỗi khi xuất tệp báo cáo: {str(e)}")
            
            # Bảng tổng hợp số liệu trích xuất chuẩn hóa
            st.markdown("#### 📊 Bảng số liệu trích xuất chuẩn hóa:")
            kpi_df = pd.DataFrame(
                [{"Chỉ số": k, "Giá trị trích xuất": str(v) if v is not None else ""} for k, v in st.session_state.final_kpi.items()]
            )
            st.dataframe(kpi_df, width="stretch")
